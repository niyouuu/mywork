import datetime
import os
import random
import docx
import qtawesome
from PyQt5 import QtWidgets, QtGui
from PyQt5.QtCore import *
from PyQt5.QtMultimedia import QMediaContent, QMediaPlayer, QAudioFormat, QAudioInput, QAudioDeviceInfo
from PyQt5.QtMultimediaWidgets import QVideoWidget
from PyQt5.QtSql import *
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
# from PyQt5.QtWebEngineWidgets import QWebEngineView
# from PyQt5.QtPrintSupport import *
import sys, sqlite3, time
from PyQt5.uic.properties import QtCore
from pydub import AudioSegment

from fullPlayDialog import fullScreenVedioDialog
from av_test import VideoRecorder
from htmlparser import MyHTMLParser


# 训练录制相关
class RecordingThread(QThread):  # 相机逻辑
    def __init__(self, duration, trainfile):
        super().__init__()
        self.duration = int(duration)
        self.trainfile = trainfile
        self.vr = VideoRecorder()

    def run(self):
        self.vr.start_recording(self.duration, self.trainfile)

    # 手动
    def set_vr_flag(self, value):
        self.vr.setflag(value)

    def set_parameter(self, duration, trainfile):
        self.duration = int(duration)
        self.trainfile = trainfile

    def set_close(self):
        self.vr.close()


class WordWindow(QWidget):
    def __init__(self, *args, **kwargs):
        super(WordWindow, self).__init__(*args, **kwargs)
        self.resize(700, 500)
        desktop = QApplication.desktop()
        self.setFixedWidth(desktop.width())
        self.setWindowTitle("欢迎使用康复训练系统")
        sc = 0
        self.sc = 0
        self.rt = 0
        self.nt = 0
        self.wt = 0
        # 播放器
        self.videowidget = QVideoWidget()
        self.player = QMediaPlayer()
        self.player.stateChanged.connect(self.on_state_changed)  # 训练录制相关
        self.audio_player = QMediaPlayer()
        self.fsDialog = fullScreenVedioDialog()
        self.media_path = ""
        # 当前用户名字
        self.temp_username = ""
        # 当前用户编号
        self.temp_userno = ""
        # 当前患者编号
        self.temp_pano = "0"
        # 训练词语数
        self.word_total = 10
        self.current_w = 0
        self.wArray = []  # 词语列表
        # 当前词语编号，在AssignBtns中赋值
        self.word_no = 0
        self.trainfolder = ""  # 训练录制相关
        self.trainfile = ""
        self.playflag = 0
        self.duration = 0
        self.is_random = 0  # 随机 0为正序
        self.is_auto = 0  # 0为自动
        self.repet_no = 3
        self.sum_num = 0  # 题目总字数
        self.wrong_num = 0  # 读错的字数
        self.recording_thread = RecordingThread(0, self.trainfile)  # 相机逻辑
        self.read_no = 0  # version2 记录读的次数
        self.ques_list = []  # 总结报告
        self.score_list = []
        self.text_list = []
        self.setUpUI()
        self.source_files = []

    def setUpUI(self):
        self.conn = sqlite3.connect("database.db")
        self.c = self.conn.cursor()
        # 添加sql语句
        self.c.close()
        self.resize(700, 500)
        try:
            # 提示
            self.msgBox = QMessageBox()
            self.msgBox.setWindowTitle("提示")
            self.msgBox.setText("请对着视频跟读")
            font = QFont()
            font.setPointSize(40)
            self.msgBox.setFont(font)
            self.msgtimer = QTimer()
            self.msgtimer.setSingleShot(True)
            self.msgtimer.timeout.connect(self.msgBox.accept)
        except Exception as e:
            print(f'Error1: {e}')

        # 选择用户
        self.layout = QVBoxLayout()
        self.indexlayout = QHBoxLayout()
        self.optionPackLayout = QHBoxLayout()
        self.optionLayout = QVBoxLayout()
        self.optionLayout1 = QHBoxLayout()
        self.optionLayout2 = QHBoxLayout()
        self.contentLayout = QHBoxLayout()
        self.video_layout = QVBoxLayout()
        self.word_layout = QVBoxLayout()
        self.word_layout1 = QHBoxLayout()
        self.word_layout2 = QHBoxLayout()

        self.titlelabel = QLabel("康复训练")
        font = self.titlelabel.font()
        font.setPointSize(25)
        font.setBold(1)
        font.setFamily("黑体")
        self.titlelabel.setFont(font)
        index_btn_len = 125
        index_btn_len_2 = 150

        self.IndexBtn = QtWidgets.QPushButton("首页")
        self.IndexBtn.setObjectName('index_button')
        self.IndexBtn.move(250, 100)
        self.IndexBtn.setFixedSize(150, 50)

        self.IndexBtn.setIcon(QIcon(QPixmap("indexres/shouye.png")))

        self.trainMissionBtn = QtWidgets.QPushButton("训练任务")
        self.trainMissionBtn.setObjectName('index_button')
        self.trainMissionBtn.setObjectName('index_button')
        self.trainMissionBtn.move(250, 100)
        self.trainMissionBtn.setFixedSize(150, 50)
        self.trainMissionBtn.setStyleSheet("QPushButton{\n"
                                           "    background:rgb(81, 71, 81);\n"
                                           "    color: white;\n"
                                           "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 24px;font-family: 微软雅黑;\n"
                                           "}\n"
                                           "QPushButton:pressed{\n"
                                           "    background:black;\n"
                                           "}")
        self.trainMissionBtn.setIcon(QIcon(QPixmap("indexres/yinpin.png")))

        # self.trainMissionBtn.setFixedWidth(index_btn_len)

        self.user_manage_btn = QtWidgets.QPushButton("用户管理")
        self.user_manage_btn.setObjectName('index_button')
        self.user_manage_btn.setObjectName('index_button')
        self.user_manage_btn.move(250, 100)
        self.user_manage_btn.setFixedSize(150, 50)
        self.user_manage_btn.setIcon(QIcon(QPixmap("indexres/yonghu.png")))

        # self.user_manage_btn.setFixedWidth(index_btn_len)

        self.sys_manage_btn = QtWidgets.QPushButton("系统管理")
        self.sys_manage_btn.setObjectName('index_button')
        self.sys_manage_btn.move(250, 100)
        self.sys_manage_btn.setFixedSize(150, 50)

        self.sys_manage_btn.setObjectName('index_button')
        self.sys_manage_btn.setIcon(QIcon(QPixmap("indexres/shezhi.png")))
        # self.sys_manage_btn.setFixedWidth(index_btn_len)

        self.studyBtn = QtWidgets.QPushButton("教 程")
        self.studyBtn.setObjectName('index_button')
        self.studyBtn.move(250, 100)
        self.studyBtn.setFixedSize(150, 50)
        self.studyBtn.setIcon(QIcon(QPixmap("indexres/dianji.png")))
        #
        # self.studyBtn.setFixedWidth(index_btn_len)

        self.jumpToLabel_5 = QLabel("        ")
        self.jumpToLabel_5.setFixedWidth(50)
        self.jumpToLabel_6 = QLabel("  ")
        self.jumpToLabel_6.setFixedWidth(60)

        self.indexlayout.addWidget(self.titlelabel)
        # self.indexlayout.addWidget(self.jumpToLabel_6)
        self.indexlayout.addWidget(self.IndexBtn)
        self.indexlayout.addWidget(self.jumpToLabel_5)
        self.indexlayout.addWidget(self.trainMissionBtn)
        self.indexlayout.addWidget(self.jumpToLabel_5)
        self.indexlayout.addWidget(self.user_manage_btn)
        self.indexlayout.addWidget(self.jumpToLabel_5)
        self.indexlayout.addWidget(self.sys_manage_btn)
        self.indexlayout.addWidget(self.jumpToLabel_5)
        self.indexlayout.addWidget(self.studyBtn)

        self.exitBtn = QtWidgets.QPushButton("退出系统")
        self.exitBtn.setObjectName('index_button')
        self.exitBtn.move(250, 100)
        self.exitBtn.setFixedSize(150, 50)
        self.exitBtn.setStyleSheet("QPushButton{\n"
                                   "    background:orange;\n"
                                   "    color:white;\n"
                                   "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 24px;font-family: 微软雅黑;\n"
                                   "}\n"
                                   "QPushButton:pressed{\n"
                                   "    background:black;\n"
                                   "}")
        self.exitBtn.setIcon(QIcon(QPixmap("indexres/tuichu.png")))

        self.indexlayout.addWidget(self.jumpToLabel_5)
        # self.indexlayout.addWidget(self.exitBtn)
        # self.exitBtn.setFixedWidth(index_btn_len)
        self.indexlayout.addWidget(self.titlelabel)
        # self.indexlayout.addWidget(self.IndexBtn)
        # self.indexlayout.addWidget(self.trainMissionBtn)
        # self.indexlayout.addWidget(self.user_manage_btn)
        # self.indexlayout.addWidget(self.sys_manage_btn)
        # self.indexlayout.addWidget(self.studyBtn)
        self.indexlayout.addWidget(self.exitBtn)
        # 导航栏按钮
        self.exitBtn.clicked.connect(self.exitBtnClick)
        self.trainMissionBtn.clicked.connect(self.trainMissionBtnClick)
        self.sys_manage_btn.clicked.connect(self.sysManageBtnClicked)
        self.studyBtn.clicked.connect(self.studyBtnClicked)

        # 添加导航栏按钮功能
        self.user_manage_btn.clicked.connect(self.userManageBtnClicked)
        # 当前已选择患者

        # 选项按钮1-10
        r = range(1, 11)
        for ele in r:
            exec('self.optionBtn{} = QPushButton("{}")'.format(ele, ele))
        self.jumpToLabel_7 = QLabel("  ")
        self.jumpToLabel_7.setFixedWidth(10)

        self.optionBtn1.move(250, 100)
        self.optionBtn1.setFixedSize(150, 50)
        self.optionBtn2.move(250, 100)
        self.optionBtn2.setFixedSize(150, 50)
        self.optionBtn3.move(250, 100)
        self.optionBtn3.setFixedSize(150, 50)
        self.optionBtn4.move(250, 100)
        self.optionBtn4.setFixedSize(150, 50)
        self.optionBtn4.move(250, 100)
        self.optionBtn5.setFixedSize(150, 50)
        self.optionBtn5.move(250, 100)
        self.optionBtn6.setFixedSize(150, 50)
        self.optionBtn6.move(250, 100)
        self.optionBtn7.setFixedSize(150, 50)
        self.optionBtn7.move(250, 100)
        self.optionBtn8.setFixedSize(150, 50)
        self.optionBtn8.move(250, 100)
        self.optionBtn9.setFixedSize(150, 50)
        self.optionBtn9.move(250, 100)
        self.optionBtn10.setFixedSize(150, 50)
        self.optionBtn10.move(250, 100)

        self.optionBtn1.setFixedSize(150, 50)
        # self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout1.addWidget(self.optionBtn1)

        self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout1.addWidget(self.optionBtn2)
        self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout1.addWidget(self.optionBtn3)
        self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout1.addWidget(self.optionBtn4)
        self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout1.addWidget(self.optionBtn5)
        # self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout2.addWidget(self.optionBtn6)
        self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionLayout2.addWidget(self.optionBtn7)
        self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionLayout2.addWidget(self.optionBtn8)
        self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionLayout2.addWidget(self.optionBtn9)
        self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionLayout2.addWidget(self.optionBtn10)

        # self.optionLayout2.addWidget(self.jumpToLabel_7)
        # # self.indexlayout.addWidget(self.exitBtn)
        # # self.exitBtn.setFixedWidth(index_btn_len)
        # self.optionLayout2.addWidget(self.optionBtn1)
        # self.optionLayout2.addWidget(self.optionBtn10)

        self.leftPageBtn = QPushButton("上一题")
        self.rightPageBtn = QPushButton("下一题")
        self.leftPageBtn.clicked.connect(self.leftBtnClick)
        self.rightPageBtn.clicked.connect(self.rightBtnClick)
        # self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionLayout.addLayout(self.optionLayout1)
        self.optionLayout.addLayout(self.optionLayout2)
        # self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionPackLayout.addWidget(self.leftPageBtn)
        self.optionPackLayout.addWidget(self.jumpToLabel_7)
        self.optionPackLayout.addLayout(self.optionLayout)
        self.optionPackLayout.addWidget(self.jumpToLabel_7)
        self.optionPackLayout.addWidget(self.rightPageBtn)

        # 视频
        self.videowidget.resize(200, 200)
        self.player.setVideoOutput(self.videowidget)
        self.playLayout = QHBoxLayout()
        self.playpause = QPushButton("播放")
        self.fullplay = QPushButton("全屏播放")
        self.playLayout.addWidget(self.playpause)
        self.playLayout.addWidget(self.fullplay)
        self.video_layout.addWidget(self.videowidget)
        self.video_layout.addLayout(self.playLayout)
        self.video_layout.setStretch(0, 4)
        self.video_layout.setStretch(1, 1)

        # 词语展示板块
        self.wordPicLabel = QLabel()
        self.wordBtn = QPushButton("发音")
        self.wordBtn.clicked.connect(self.wordBtnClick)
        self.wordLabel = QLabel("")
        font = QFont()
        font.setPixelSize(50)
        font.setBold(1)
        font.setFamily("黑体")
        self.wordLabel.setFont(font)

        # 录音相关
        # self.temp_pano = "0"
        self.train_str_no = ""
        self.file_path = ""
        self.trial_no = "0"  # 尚未自动生成
        from myaudiolayout import AudioRecorder
        self.vbox = QVBoxLayout()
        self.audiorecorder = AudioRecorder(self.temp_pano, self.train_str_no, self.trial_no, parent=self)
        self.audiorecorder.pagestate_clicked.connect(self.setBtnsState)  # 改变上下翻页按钮是否激活，防止录音被打断
        self.audiorecorder.hide()  # 去除录音
        self.endRecordBtn = QPushButton("结束录制")  # 手动
        self.wordBtn.setFixedSize(150, 50)
        self.wordBtn.move(250, 100)
        self.endRecordBtn.setFixedSize(150, 50)
        self.endRecordBtn.move(250, 100)
        font.setPixelSize(25)
        self.endRecordBtn.setFont(font)
        self.endRecordBtn.clicked.connect(self.endRecordBtnClicked)
        from assess import AudioAssess  # version2
        self.audio_assess = AudioAssess()
        self.audio_assess.set_hide()
        self.vbox.addWidget(self.audiorecorder)
        self.vbox.addWidget(self.endRecordBtn,1, Qt.AlignCenter)  # 手动
        self.vbox.addWidget(self.audio_assess)

        # 总排版
        self.layout.addLayout(self.indexlayout)
        self.layout.addLayout(self.optionPackLayout)
        self.layout.addLayout(self.contentLayout)
        # 连词成句按钮板块
        self.word_layout1.addWidget(self.wordPicLabel, 1, Qt.AlignCenter)
        self.word_layout2.addWidget(self.wordBtn)
        self.word_layout.addLayout(self.word_layout1)
        self.word_layout.addWidget(self.wordLabel, 1, Qt.AlignCenter)
        self.word_layout.addLayout(self.word_layout2)
        # 录音相关
        self.word_layout.addLayout(self.vbox)

        # 连词成句按钮板块排版
        self.contentLayout.addLayout(self.video_layout, 2)
        self.contentLayout.addLayout(self.word_layout, 3)
        self.playpause.clicked.connect(self.videoBtnClick)
        self.fullplay.clicked.connect(self.fullopenVedio)

        self.setLayout(self.layout)

        lab = [self.endRecordBtn]
        tb = [self.exitBtn, self.sys_manage_btn, self.studyBtn, self.user_manage_btn, self.IndexBtn
              ]  # ,
        yb = [self.optionBtn1, self.optionBtn2, self.optionBtn3, self.optionBtn4, self.optionBtn5, self.optionBtn6,
              self.optionBtn7, self.optionBtn8, self.optionBtn9, self.optionBtn10,
              ]
        lb = [self.leftPageBtn, self.rightPageBtn]
        hb = [self.playpause, self.fullplay]
        qb = [self.wordBtn]
        font = QtGui.QFont()
        font.setPointSize(15)  # 括号里的数字可以设置成自己想要的字体大小
        # font.setFamily("SimHei")  # 黑体
        font.setFamily("SimSun")  # 宋体
        for i in qb:
            i.setFixedSize(1000, 35)
            i.setStyleSheet("QPushButton{\n"
                            "    background-color: rgb(170, 255, 127);\n"
                            "    color: rgb(81, 71, 81);\n"
                            "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 24px;font-family: 微软雅黑;\n"
                            "}\n"
                            "QPushButton:pressed{\n"
                            "    background:yellow;\n"
                            "}")
            i.setFont(font)
            # font = QtGui.QFont()
            # font.setPointSize(10)  # 括号里的数字可以设置成自己想要的字体大小
            # i.setFont(font)
        for i in lab:
            i.setFixedSize(1000, 35)
            i.setStyleSheet("QPushButton{\n"
                            "    background-color: rgb(193, 193, 193);\n"
                            "    color: rgb(81, 71, 81);\n"
                            "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 24px;font-family: 微软雅黑;\n"
                            "}\n"
                            "QPushButton:pressed{\n"
                            "    background:black;\n"
                            "}")
            i.setFont(font)
        for i in tb:
            i.setStyleSheet("QPushButton{\n"
                            "    background:rgb(244, 183, 0);\n"
                            "    color: rgb(81, 71, 81);\n"
                            "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 24px;font-family: 微软雅黑;\n"
                            "}\n"
                            "QPushButton:pressed{\n"
                            "    background:black;\n"
                            "}")
            i.setFont(font)
        for i in yb:
            i.setFixedSize(150, 30)
            i.setStyleSheet("QPushButton{\n"
                            "    background-color: rgb(170, 255, 127);\n"
                            "    color: rgb(81, 71, 81);\n"
                            "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 5px;font-family: 微软雅黑;\n"
                            "}\n"
                            "QPushButton:pressed{\n"
                            "    background:yellow;\n"
                            "}")
            i.setFont(font)
            font = QtGui.QFont()
            font.setPointSize(10)  # 括号里的数字可以设置成自己想要的字体大小
            i.setFont(font)
        for i in lb:
            i.setFixedSize(450, 35)
            i.setStyleSheet("QPushButton{\n"
                            "    background-color: rgb(170, 255, 127);\n"
                            "    color: rgb(81, 71, 81);\n"
                            "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 5px;font-family: 微软雅黑;\n"
                            "}\n"
                            "QPushButton:pressed{\n"
                            "    background:yellow;\n"
                            "}")
            i.setFont(font)
            font = QtGui.QFont()
            font.setPointSize(10)  # 括号里的数字可以设置成自己想要的字体大小
            i.setFont(font)
        for i in hb:
            i.setFixedSize(400, 35)
            i.setStyleSheet("QPushButton{\n"
                            "    background-color: rgb(170, 255, 127);\n"
                            "    color: rgb(81, 71, 81);\n"
                            "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 5px;font-family: 微软雅黑;\n"
                            "}\n"
                            "QPushButton:pressed{\n"
                            "    background:yellow;\n"
                            "}")
            i.setFont(font)
            font = QtGui.QFont()
            font.setPointSize(10)  # 括号里的数字可以设置成自己想要的字体大小
            i.setFont(font)

    # 录音相关，在录音时控制翻页是否可用
    def setBtnsState(self, b_value):
        self.leftPageBtn.setEnabled(b_value)
        self.rightPageBtn.setEnabled(b_value)
        self.endRecordBtn.setEnabled(not b_value)

    def IndexBtnClick(self):
        from t_main import IndexWindow
        self.mainwindow = IndexWindow()
        self.mainwindow.showFullScreen()
        self.mainwindow.searchButtonClicked()
        self.close()

    def trainMissionBtnClick(self):
        if not (self.temp_pano == ""):
            from train import TrainWindow
            self.trainWindow = TrainWindow()
            self.trainWindow.showFullScreen()
            self.trainWindow.setNo(self.temp_pano)
            self.close()
        else:
            print(QMessageBox.warning(self, "警告", "请选择一名患者", QMessageBox.Yes, QMessageBox.Yes))

    def wordBtnClick(self):
        self.audio_player.setVolume(80)
        self.audio_player.play()
        return

    def AssignWord(self):
        db = QSqlDatabase.addDatabase("QSQLITE")
        db.setDatabaseName('database.db')
        db.open()
        query = QSqlQuery()
        sql = "SELECT COUNT(*) FROM words"
        query.exec_(sql)
        query.next()
        word_in_sql = int(query.value(0))
        if (word_in_sql < self.word_total):
            temp_num = word_in_sql
            self.word_total = temp_num
        else:
            temp_num = self.word_total
        if self.is_random == 0:
            sql = "SELECT * FROM words ORDER BY RANDOM() LIMIT '%s'" % (temp_num)
        else:
            sql = "SELECT * FROM words ORDER BY RANDOM() LIMIT '%s'" % (temp_num)  # 随机
        query.exec_(sql)
        r = range(0, temp_num)
        for ele in r:
            query.next()
            self.wArray.append([])
            print(query.value(0))
            self.wArray[ele].append(query.value(0))  # word_no
            self.wArray[ele].append(query.value(1))  # word_name
            self.wArray[ele].append(query.value(2))  # word_audio
            self.wArray[ele].append(query.value(3))  # word_image
        # 训练录制相关
        # 定义查询语句和插入语句
        select_sql = "SELECT trial_no FROM trial WHERE pa_no = ?"
        insert_sql = "INSERT INTO trial (pa_no, trial_no) VALUES (?, 0)"
        update_sql = "UPDATE trial SET trial_no = ? WHERE pa_no = ?"
        # 定义要查询的pa_no值
        pa_no = self.temp_pano
        # 执行查询操作
        query = QSqlQuery()
        query.prepare(select_sql)
        query.addBindValue(pa_no)
        query.exec_()
        if not query.next():
            # 如果查询结果为空，则插入新记录
            query.prepare(insert_sql)
            query.addBindValue(pa_no)
            query.exec_()
            db.commit()
            self.trial_no = 0
        else:
            # 如果查询结果不为空，则更新记录
            trial_no = query.value(0) + 1
            query.prepare(update_sql)
            query.addBindValue(trial_no)
            query.addBindValue(pa_no)
            query.exec_()
            db.commit()
            self.trial_no = trial_no

    def AssignBtns(self):
        # 录音相关
        self.word_no = self.wArray[self.current_w][0]
        self.train_str_no = "w" + str(self.word_no)
        self.audiorecorder.setTrainStr(self.train_str_no)
        self.audiorecorder.setTrialNo(self.trial_no)
        self.audiorecorder.setFilePath(self.trial_no, self.train_str_no)
        self.audiorecorder.setAuBtnsState(False)
        # 设置词语按钮
        word_name = self.wArray[self.current_w][1]
        image_path = self.wArray[self.current_w][3]
        self.wordLabel.setText(word_name)
        temp_pic = QtGui.QPixmap(image_path)
        w = temp_pic.width()
        h = temp_pic.height()
        if not w == 0:
            if (w / h > 1):
                temp_pic = temp_pic.scaledToWidth(700)
            else:
                temp_pic = temp_pic.scaledToHeight(600)
        # temp_pic = QtGui.QPixmap(image_path).scaledToWidth(700)
        self.wordPicLabel.setPixmap(temp_pic)
        self.wordPicLabel.setFixedWidth(700)
        self.wordPicLabel.setScaledContents(True)
        # 设置播放
        self.media_path = "pa_video/pa" + str(self.temp_pano) + "/pa" + str(self.temp_pano) + "_w" + str(
            self.wArray[self.current_w][0]) + ".mp4"
        self.player.setMedia(QMediaContent(QUrl.fromLocalFile(self.media_path)))
        audio_path = "pa_audio/pa" + str(self.temp_pano) + "/word/" + str(self.word_no) + ".mp3"  # 训练录制相关
        if not os.path.exists(audio_path):
            audio_path = str(self.wArray[self.current_w][2])
        self.audio_player.setMedia(QMediaContent(QUrl.fromLocalFile(audio_path)))
        self.videoBtnClick()

    def videoBtnClick(self):
        self.player.setVolume(80)
        self.player.play()
        return

    def fullopenVedio(self):
        self.fsDialog.setMedia(self.media_path)
        self.fsDialog.showFullScreen()
        self.fsDialog.play()
        return

    def leftBtnClick(self):
        if (self.current_w > 0):
            self.current_w -= 1
            self.AssignBtns()

    def rightBtnClick(self):
        self.playflag = 0  # 训练录制相关
        self.read_no = 0  # version2
        if (self.current_w < self.word_total - 1):
            self.current_w += 1
            self.AssignBtns()
        else:
            self.recording_thread.set_close()  # 相机逻辑
            self.writelogs()
            # self.wtt += self.wt
            # self.rtt += self.rt
            # self.ntt += self.nt
            ress = QMessageBox.information(self, "提示", "是否查看本次训练结果?", QMessageBox.Yes, QMessageBox.No)
            if (ress == QMessageBox.Yes):
                self.writetubiao()
                self.merge_html_files()
            else:
                return
            # 总结报告
            from train_log import TrainDialogue
            traindialogue = TrainDialogue(self)
            traindialogue.exec_()
            self.trainMissionBtnClick()
    def writetubiao(self):
        from pyecharts import options as opts
        from pyecharts.charts import Bar
        now = datetime.datetime.now()
        # 将日期格式化为SQLite可接受的格式
        log_date = now.strftime('%Y-%m-%d')
        c = (
            Bar()
            .add_xaxis(
                [
                    log_date
                ]
            )
            .add_yaxis("错题数", [self.wt])
            .add_yaxis("普通读音数", [self.nt])
            .add_yaxis("正确数", [self.rt])
            .set_global_opts(
                # xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=-15)),  # 此行代码旋转标签的角度
                title_opts=opts.TitleOpts(title="患者"+self.temp_pano +"词语训练"),
            )
            .render(self.temp_pano+"word.html")
        )

        html_file_path = self.temp_pano + "word.html"
        self.source_files.append(html_file_path)

        import webbrowser
        webbrowser.open(self.temp_pano+'word.html')

    def merge_html_files(self):
        target_file_path = self.temp_pano + 'target_merged.html'  # 假设target_merged.html应该在工作目录中

        # 检查target_merged.html文件是否存在，如果不存在则创建
        if not os.path.exists(target_file_path):
            with open(target_file_path, 'w', encoding='utf-8') as file:
                file.write('<html><body></body></html>')  # 创建一个简单的HTML结构

        # 读取target_merged.html文件的内容
        with open(target_file_path, 'r', encoding='utf-8') as file:
            target_content = file.read()

        # 在target_merged.html文件中找到合适的位置插入源HTML文件的内容
        insert_position = target_content.find('</body>')
        if insert_position != -1:
            # 插入所有源HTML文件的内容
            for source_file in self.source_files:
                if os.path.exists(source_file):
                    with open(source_file, 'r', encoding='utf-8') as file:
                        source_content = file.read()
                        target_content = target_content[:insert_position] + source_content + target_content[
                                                                                             insert_position:]
                else:
                    print(f"File {source_file} does not exist.")
        else:
            # 如果没有找到</body>标签，可以选择在末尾添加或者进行其他操作
            for source_file in self.source_files:
                if os.path.exists(source_file):
                    with open(source_file, 'r', encoding='utf-8') as file:
                        source_content = file.read()
                        target_content += source_content
                else:
                    print(f"File {source_file} does not exist.")

        # 将新的内容写回target_merged.html文件
        with open(target_file_path, 'w', encoding='utf-8') as file:
            file.write(target_content)

    def writealltubiao(self):
        from pyecharts import options as opts
        from pyecharts.charts import Bar

        now = datetime.datetime.now()
        # 将日期格式化为SQLite可接受的格式
        log_date = now.strftime('%Y-%m-%d')
        c = (
            Bar()
            .add_xaxis(
                [
                    "第" + self.trial_no + "次训练"
                ]
            )
            .add_yaxis("分数", [self.sc])
            .set_global_opts(
                # xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=-15)),  # 此行代码旋转标签的角度
                title_opts=opts.TitleOpts(title="患者" + self.temp_pano,
                                          subtitle=log_date),
            )
            .render("label.html")
        )
        import webbrowser
        webbrowser.open('label.html')

    def writelogs(self):
        # 连接到数据库
        conn = sqlite3.connect('database.db')
        # 创建游标
        cursor = conn.cursor()
        # 查询是否存在匹配的记录
        query = '''
            SELECT * FROM logs WHERE pa_no = ? AND log_date = ?
        '''
        # 获取当前日期
        now = datetime.datetime.now()
        # 将日期格式化为SQLite可接受的格式
        log_date = now.strftime('%Y-%m-%d')
        cursor.execute(query, (self.temp_pano, log_date))
        record = cursor.fetchone()
        # 如果不存在记录，则插入新记录
        if not record:
            # 计算acc的值
            acc = 100 - int(self.wrong_num / self.sum_num * 100)
            # 插入新记录
            query = '''
                INSERT INTO logs (pa_no, log_date, wrong_num, sum_num, acc)
                VALUES (?, ?, ?, ?, ?)
            '''
            values = (self.temp_pano, log_date, self.wrong_num, self.sum_num, acc)
            cursor.execute(query, values)
        else:
            # 计算更新后的wrong_num、sum_num和acc的值
            wrong_num = record[3] + self.wrong_num
            sum_num = record[4] + self.sum_num
            acc = 100 - int(wrong_num / sum_num * 100)
            # 更新记录
            query = '''
                UPDATE logs SET wrong_num = ?, sum_num = ?, acc = ?
                WHERE pa_no = ? AND log_date = ?
            '''
            values = (wrong_num, sum_num, acc, self.temp_pano, log_date)
            cursor.execute(query, values)
        # 提交更改并关闭游标和连接
        conn.commit()
        cursor.close()
        conn.close()
        # 创建docx文档
        doc = docx.Document()
        p = doc.add_paragraph()
        p.add_run(log_date + '\n')
        # 添加窗口文本到文档中
        for i, (ques, score, text) in enumerate(zip(self.ques_list, self.score_list, self.text_list)):
            # 将窗口文本添加到docx文档中
            p.add_run(ques + '\n').bold = True  # 加粗显示问题
            p.add_run(score + '\n')  # 显示分数
            # 使用 MyHTMLParser 解析 HTML 文本
            parser = MyHTMLParser()
            parser.feed(text)
            # 处理文本和属性
            for parsed_text, bold, color in parser.result:
                run = p.add_run(parsed_text)
                # 处理加粗
                if bold:
                    run.bold = True
                # 处理颜色
                if color:
                    run.font.color.rgb = docx.shared.RGBColor(*color)
        filename = 'pa_train/pa_' + str(self.temp_pano) + '/trial' + str(self.trial_no) + '/output.docx'
        # 保存为docx文件
        doc.save(filename)

        filename = 'pa_train/pa_' + str(self.temp_pano) + '/output.docx'
        if os.path.exists(filename):
            # 如果文件已经存在，则打开现有文件以追加模式
            doc = docx.Document(filename)
        else:
            # 如果文件不存在，则创建一个新文件
            doc = docx.Document()
        p = doc.add_paragraph()
        p.add_run(log_date + '\n')
        p.add_run('trial' + str(self.trial_no) + '\n')
        # 在现有文档或新文档中添加窗口文本
        for i, (ques, score, text) in enumerate(zip(self.ques_list, self.score_list, self.text_list)):
            # 将窗口文本添加到docx文档中
            p = doc.add_paragraph()
            p.add_run(ques + '\n').bold = True  # 加粗显示问题
            p.add_run(score + '\n')  # 显示分数
            # 使用 MyHTMLParser 解析 HTML 文本
            parser = MyHTMLParser()
            parser.feed(text)
            # 处理文本和属性
            for parsed_text, bold, color in parser.result:
                run = p.add_run(parsed_text)
                # 处理加粗
                if bold:
                    run.bold = True
                # 处理颜色
                if color:
                    run.font.color.rgb = docx.shared.RGBColor(*color)
        # 保存为docx文件
        doc.save(filename)

    def exitBtnClick(self):
        ret = QMessageBox.information(self, "提示", "是否退出系统?", QMessageBox.Yes, QMessageBox.No)
        if (ret == QMessageBox.Yes):
            sys.exit(app.exec_())
        else:
            return

    def sysManageBtnClicked(self):
        from sys_management import SysManageWindow
        self.sysManageWindow = SysManageWindow()
        self.sysManageWindow.searchButtonClicked()
        self.sysManageWindow.showFullScreen()
        self.close()

    def setTaskNum(self, pano_value, num_value, is_random, is_auto, repet_no):
        self.temp_pano = pano_value
        self.word_total = num_value
        self.is_random = is_random
        self.is_auto = is_auto
        if repet_no >= 1:
            self.repet_no = repet_no

    def setTrainFile(self):  # 训练录制相关
        self.trainfolder = "pa_train/pa_" + str(self.temp_pano) + "/trial" + str(self.trial_no) + "/"
        if not os.path.exists(self.trainfolder):
            os.makedirs(self.trainfolder)
        self.trainfile = self.trainfolder + self.train_str_no + "_" + str(self.read_no) + ".mp4"  # version2

    def on_state_changed(self, state):
        try:
            if self.is_auto == 0:
                if state == QMediaPlayer.StoppedState and self.playflag <= self.repet_no:
                    self.read_no = self.read_no + 1  # version2
                    self.duration = (self.player.duration() / 1000 * 2 + 5)
                    # 训练录制相关
                    self.setTrainFile()
                    # print("录制中" + str(self.current_w))
                    self.recording_thread.set_parameter(self.duration, self.trainfile)  # 相机逻辑
                    self.recording_thread.start()
                    self.setBtnsState(False)
                    self.msgBox.setText("请跟读!!")
                    self.msgtimer.start(700)  # 提示
                    self.msgBox.exec_()
                    self.playflag = self.playflag + 1
                    time.sleep(self.duration)
                    temp_standard_text = self.wArray[self.current_w][1]
                    temp_standard_audio = "audio/word/" + str(self.word_no) + ".mp3"
                    temp_user_audio = "user_audio.mp3"
                    # 总结报告
                    score, text, wrong_num, sum_num = self.audio_assess.set_file(temp_standard_text,
                                                                                 temp_standard_audio, temp_user_audio)
                    self.sum_num += sum_num
                    self.wrong_num += wrong_num
                    os.remove("user_audio.mp3")
                    ques = "第" + str(self.read_no) + "遍：" + temp_standard_text
                    self.record_score_and_text(ques, score, text)
                    self.sc = self.audio_assess.sc
                    # print("这是词语界面的" + self.sc)
                    # from assess import AudioAssess
                    if int(self.sc) > 85:
                        self.rt += 1
                    elif int(self.sc) <= 85 and int(self.sc) >= 70:
                        self.nt += 1
                    else:
                        self.wt += 1

                    time.sleep(1)
                    if self.playflag == self.repet_no:
                        self.rightBtnClick()
                    else:
                        self.videoBtnClick()
            # 手动
            elif self.is_auto == 1:
                if state == QMediaPlayer.StoppedState:
                    self.read_no = self.read_no + 1  # version2
                    # 训练录制相关
                    self.setTrainFile()
                    print("录制中" + str(self.current_w))
                    self.duration = 0
                    self.recording_thread.set_parameter(self.duration, self.trainfile)  # 相机逻辑
                    self.recording_thread.start()
                    self.setBtnsState(False)
                    self.msgBox.setText("按 结束录制 结束跟读")
                    self.msgtimer.start(700)  # 提示
                    self.sc = self.audio_assess.sc
                    print("这是元音界面的" + self.sc)
                    from assess import AudioAssess
                    if int(self.sc) > 85:
                        self.rt += 1
                    elif int(self.sc) <= 85 and int(self.sc) >= 70:
                        self.nt += 1
                    else:
                        self.wt += 1
                    self.msgBox.exec_()
                    self.playflag = self.playflag + 1
        except Exception as e:
            print(f'Error2: {e}')

    # 手动
    def endRecordBtnClicked(self):
        try:
            self.recording_thread.set_vr_flag(True)
            time.sleep(1)
            temp_standard_audio = "audio/word/" + str(self.word_no) + ".mp3"
            temp_user_audio = "user_audio.mp3"
            temp_standard_text = self.wArray[self.current_w][1]
            # 总结报告
            score, text, wrong_num, sum_num = self.audio_assess.set_file(temp_standard_text, temp_standard_audio,
                                                                         temp_user_audio)
            self.wrong_num += wrong_num
            self.sum_num += sum_num
            os.remove("user_audio.mp3")
            ques = "第" + str(self.read_no) + "遍：" + temp_standard_text
            self.record_score_and_text(ques, score, text)
            time.sleep(1)
            self.setBtnsState(True)
            self.recording_thread.set_vr_flag(False)
        except Exception as e:
            print(f'Error3: {e}')

    # 总结报告
    def record_score_and_text(self, ques, score, text):
        self.ques_list.append(ques)
        self.score_list.append(score)
        self.text_list.append(text)

    def userManageBtnClicked(self):
        self.userManageWindow.showFullScreen()
        self.userManageWindow.searchButtonClicked()
        self.close()

    def studyBtnClicked(self):
        QtGui.QDesktopServices.openUrl(QtCore.QUrl('https://book.yunzhan365.com/vdogo/mxvi/mobile/index.html'))


if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon("./images/MainWindow_1.png"))
    temp_pano = 0
    mainWindow = WordWindow()
    mainWindow.setTaskNum(temp_pano, int(5), is_random=1, is_auto=1, repet_no=1)
    mainWindow.AssignWord()
    mainWindow.AssignBtns()
    mainWindow.showFullScreen()
    # mainWindow.show()
    sys.exit(app.exec_())