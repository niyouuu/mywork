import datetime
import os
import random

import docx
import qtawesome
from PyQt5 import QtWidgets, QtGui
from PyQt5.QtCore import *
from PyQt5.QtMultimedia import QMediaContent, QMediaPlayer
from PyQt5.QtMultimediaWidgets import QVideoWidget
from PyQt5.QtSql import *
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
#from PyQt5.QtWebEngineWidgets import QWebEngineView
#from PyQt5.QtPrintSupport import *
import sys,sqlite3,time
from PyQt5.uic.properties import QtCore
from fullPlayDialog import fullScreenVedioDialog
from av_test import VideoRecorder
from htmlparser import MyHTMLParser


from PyQt5.QtCore import QObject, QTimer, pyqtSignal

class TimerThread(QObject):
    timer_signal = pyqtSignal()  # define a signal to emit when the timer is up

    def __init__(self, interval):
        super().__init__()
        self.interval = int(interval)
        self.is_running = False

    def start(self):
        self.is_running = True
        QTimer.singleShot(self.interval * 1000, self.timeout)  # QTimer uses milliseconds, so multiply by 1000

    def timeout(self):
        print("out")
        if self.is_running:
            self.timer_signal.emit()
            self.is_running = False

    def stop(self):
        self.is_running = False

    def change_interval(self, interval):
        self.interval = interval
        self.is_running = True
        QTimer.singleShot(self.interval * 1000, self.timeout)

#训练录制相关
class RecordingThread(QThread):# 相机逻辑
    def __init__(self, duration, trainfile):
        super().__init__()
        self.duration = int(duration)
        self.trainfile = trainfile
        self.vr = VideoRecorder()

    def run(self):
        self.vr.start_recording(self.duration, self.trainfile)
    # 手动
    def set_vr_flag(self,value):
        self.vr.setflag(value)

    def set_parameter(self, duration, trainfile):
        self.duration = int(duration)
        self.trainfile = trainfile

    def set_close(self):
        self.vr.close()

class WSWindow(QWidget):
    def __init__(self, *args, **kwargs):
        super(WSWindow, self).__init__(*args, **kwargs)
        self.resize(700, 500)
        icon = QIcon()
        icon.addPixmap(QPixmap('logo.jpg'))
        self.setWindowTitle("欢迎使用康复训练系统")
        # 播放器
        sc = 0
        self.sc = 0
        self.wrongword = 0
        self.rightword = 0
        self.rt = 0
        self.nt = 0
        self.wt = 0
        self.videowidget = QVideoWidget()
        self.player = QMediaPlayer()
        self.player.stateChanged.connect(self.on_state_changed)  # 训练录制相关
        self.audio_player = QMediaPlayer()
        self.fsDialog = fullScreenVedioDialog()
        self.media_path = ""
        # 当前用户名字
        self.temp_username = ""
        # 当前用户编号
        self.temp_userno = ""
        # 当前患者编号
        self.temp_pano = "0"
        self.trainfolder = ""  # 训练录制相关
        self.trainfile = ""
        self.playflag = 0
        self.duration = 0
        self.timer_thread = TimerThread(interval=10)
        self.timer_thread.timer_signal.connect(self.timer_callback)
        self.time_flag=0
        self.is_random = 0  # 随机 0为正序
        self.is_auto = 0  # 0为自动
        self.repet_no=3
        self.sum_num = 0  # 题目总字数
        self.wrong_num = 0  # 读错的字数
        self.recording_thread = RecordingThread(0, self.trainfile)  # 相机逻辑
        self.read_no=0 #version2 记录读的次数
        self.ques_list = []# 总结报告
        self.score_list = []
        self.text_list = []
        # 连词计数
        self.i = 1
        self.countflag = 0
        # 训练句子数
        self.sentence_total = 5
        self.sArray = []  # 句子列表
        self.current_s = 0
        self.wArray = []  # 词语列表
        self.w_btnArray = []  # 按钮表
        self.s_btnArray = []  # 句子按钮表
        # 句子按钮
        self.s_btn1 = QPushButton("")
        self.s_btn2 = QPushButton("")
        self.s_btn3 = QPushButton("")
        self.s_btn4 = QPushButton("")
        self.s_btn5 = QPushButton("")
        self.s_btn6 = QPushButton("")
        self.s_btn7 = QPushButton("")
        self.s_btn8 = QPushButton("")
        self.s_btn9 = QPushButton("")
        self.s_btn10 = QPushButton("")
        self.setUpUI()

    def setUpUI(self):
        self.conn = sqlite3.connect("database.db")
        self.c = self.conn.cursor()
        # 添加sql语句
        self.c.close()
        self.resize(700,500)
        try:
            # 提示
            self.msgBox = QMessageBox()
            self.msgBox.setWindowTitle("提示")
            self.msgBox.setText("请对着视频跟读")
            font = QFont()
            font.setPointSize(40)
            self.msgBox.setFont(font)
            self.msgtimer = QTimer()
            self.msgtimer.setSingleShot(True)
            self.msgtimer.timeout.connect(self.msgBox.accept)
        except Exception as e:
            print(f'Error: {e}')

        # 选择用户
        self.layout = QVBoxLayout()
        self.indexlayout = QHBoxLayout()
        self.optionPackLayout = QHBoxLayout()
        self.optionLayout = QVBoxLayout()
        self.optionLayout1 = QHBoxLayout()
        self.optionLayout2 = QHBoxLayout()
        self.contentLayout = QHBoxLayout()
        self.video_layout = QVBoxLayout()
        self.ws_layout = QVBoxLayout()
        self.ws_layout1 = QHBoxLayout()
        self.ws_layout2 = QHBoxLayout()
        self.ws_layout3 = QHBoxLayout()
        self.ws_layout4 = QHBoxLayout()

        self.titlelabel = QLabel("康复训练")
        font = self.titlelabel.font()
        font.setPointSize(25)
        font.setBold(1)
        font.setFamily("黑体")
        self.titlelabel.setFont(font)
        index_btn_len = 125
        index_btn_len_2 = 150

        self.IndexBtn = QtWidgets.QPushButton("首页")
        self.IndexBtn.setObjectName('index_button')
        self.IndexBtn.move(250, 100)
        self.IndexBtn.setFixedSize(150, 50)

        self.IndexBtn.setIcon(QIcon(QPixmap("indexres/shouye.png")))

        self.trainMissionBtn = QtWidgets.QPushButton("训练任务")
        self.trainMissionBtn.setObjectName('index_button')
        self.trainMissionBtn.setObjectName('index_button')
        self.trainMissionBtn.move(250, 100)
        self.trainMissionBtn.setFixedSize(150, 50)
        self.trainMissionBtn.setStyleSheet("QPushButton{\n"
                                           "    background:rgb(81, 71, 81);\n"
                                           "    color: white;\n"
                                           "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 24px;font-family: 微软雅黑;\n"
                                           "}\n"
                                           "QPushButton:pressed{\n"
                                           "    background:black;\n"
                                           "}")
        self.trainMissionBtn.setIcon(QIcon(QPixmap("indexres/yinpin.png")))

        # self.trainMissionBtn.setFixedWidth(index_btn_len)

        self.user_manage_btn = QtWidgets.QPushButton("用户管理")
        self.user_manage_btn.setObjectName('index_button')
        self.user_manage_btn.setObjectName('index_button')
        self.user_manage_btn.move(250, 100)
        self.user_manage_btn.setFixedSize(150, 50)
        self.user_manage_btn.setIcon(QIcon(QPixmap("indexres/yonghu.png")))

        # self.user_manage_btn.setFixedWidth(index_btn_len)

        self.sys_manage_btn = QtWidgets.QPushButton("系统管理")
        self.sys_manage_btn.setObjectName('index_button')
        self.sys_manage_btn.move(250, 100)
        self.sys_manage_btn.setFixedSize(150, 50)

        self.sys_manage_btn.setObjectName('index_button')
        self.sys_manage_btn.setIcon(QIcon(QPixmap("indexres/shezhi.png")))
        # self.sys_manage_btn.setFixedWidth(index_btn_len)

        self.studyBtn = QtWidgets.QPushButton("教 程")
        self.studyBtn.setObjectName('index_button')
        self.studyBtn.move(250, 100)
        self.studyBtn.setFixedSize(150, 50)
        self.studyBtn.setIcon(QIcon(QPixmap("indexres/dianji.png")))
        #
        # self.studyBtn.setFixedWidth(index_btn_len)

        self.jumpToLabel_5 = QLabel("        ")
        self.jumpToLabel_5.setFixedWidth(50)
        self.jumpToLabel_6 = QLabel("  ")
        self.jumpToLabel_6.setFixedWidth(60)

        self.indexlayout.addWidget(self.titlelabel)
        # self.indexlayout.addWidget(self.jumpToLabel_6)
        self.indexlayout.addWidget(self.IndexBtn)
        self.indexlayout.addWidget(self.jumpToLabel_5)
        self.indexlayout.addWidget(self.trainMissionBtn)
        self.indexlayout.addWidget(self.jumpToLabel_5)
        self.indexlayout.addWidget(self.user_manage_btn)
        self.indexlayout.addWidget(self.jumpToLabel_5)
        self.indexlayout.addWidget(self.sys_manage_btn)
        self.indexlayout.addWidget(self.jumpToLabel_5)
        self.indexlayout.addWidget(self.studyBtn)

        self.exitBtn = QtWidgets.QPushButton("退出系统")
        self.exitBtn.setObjectName('index_button')
        self.exitBtn.move(250, 100)
        self.exitBtn.setFixedSize(150, 50)
        self.exitBtn.setStyleSheet("QPushButton{\n"
                                   "    background:orange;\n"
                                   "    color:white;\n"
                                   "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 24px;font-family: 微软雅黑;\n"
                                   "}\n"
                                   "QPushButton:pressed{\n"
                                   "    background:black;\n"
                                   "}")
        self.exitBtn.setIcon(QIcon(QPixmap("indexres/tuichu.png")))

        self.indexlayout.addWidget(self.jumpToLabel_5)
        # self.indexlayout.addWidget(self.exitBtn)
        # self.exitBtn.setFixedWidth(index_btn_len)
        self.indexlayout.addWidget(self.titlelabel)
        # self.indexlayout.addWidget(self.IndexBtn)
        # self.indexlayout.addWidget(self.trainMissionBtn)
        # self.indexlayout.addWidget(self.user_manage_btn)
        # self.indexlayout.addWidget(self.sys_manage_btn)
        # self.indexlayout.addWidget(self.studyBtn)
        self.indexlayout.addWidget(self.exitBtn)
        # 导航栏按钮
        self.exitBtn.clicked.connect(self.exitBtnClick)
        self.trainMissionBtn.clicked.connect(self.trainMissionBtnClick)
        self.sys_manage_btn.clicked.connect(self.sysManageBtnClicked)
        self.studyBtn.clicked.connect(self.studyBtnClicked)

        # 添加导航栏按钮功能
        self.user_manage_btn.clicked.connect(self.userManageBtnClicked)
        # 当前已选择患者

        # 选项按钮1-10
        r = range(1, 11)
        for ele in r:
            exec('self.optionBtn{} = QPushButton("{}")'.format(ele, ele))
        self.jumpToLabel_7 = QLabel("  ")
        self.jumpToLabel_7.setFixedWidth(10)

        self.optionBtn1.move(250, 100)
        self.optionBtn1.setFixedSize(150, 50)
        self.optionBtn2.move(250, 100)
        self.optionBtn2.setFixedSize(150, 50)
        self.optionBtn3.move(250, 100)
        self.optionBtn3.setFixedSize(150, 50)
        self.optionBtn4.move(250, 100)
        self.optionBtn4.setFixedSize(150, 50)
        self.optionBtn4.move(250, 100)
        self.optionBtn5.setFixedSize(150, 50)
        self.optionBtn5.move(250, 100)
        self.optionBtn6.setFixedSize(150, 50)
        self.optionBtn6.move(250, 100)
        self.optionBtn7.setFixedSize(150, 50)
        self.optionBtn7.move(250, 100)
        self.optionBtn8.setFixedSize(150, 50)
        self.optionBtn8.move(250, 100)
        self.optionBtn9.setFixedSize(150, 50)
        self.optionBtn9.move(250, 100)
        self.optionBtn10.setFixedSize(150, 50)
        self.optionBtn10.move(250, 100)

        self.optionBtn1.setFixedSize(150, 50)
        # self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout1.addWidget(self.optionBtn1)

        self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout1.addWidget(self.optionBtn2)
        self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout1.addWidget(self.optionBtn3)
        self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout1.addWidget(self.optionBtn4)
        self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout1.addWidget(self.optionBtn5)
        # self.optionLayout1.addWidget(self.jumpToLabel_7)
        self.optionLayout2.addWidget(self.optionBtn6)
        self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionLayout2.addWidget(self.optionBtn7)
        self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionLayout2.addWidget(self.optionBtn8)
        self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionLayout2.addWidget(self.optionBtn9)
        self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionLayout2.addWidget(self.optionBtn10)

        self.timuBtn = QPushButton("跳转")
        self.timuBtn.clicked.connect(self.timuBtnClick)
        self.jumpToLabel = QLabel("跳转到第")
        self.jumpToLabel_2 = QLabel("   ")
        self.jumpToLabel_4 = QLabel("   ")
        self.pageEdit = QLineEdit()
        self.pageEdit.setFixedWidth(100)
        s =  "题"
        self.pageLabel = QLabel(s)
        Hlayout = QHBoxLayout()
        Hlayout2 = QHBoxLayout
        Hlayout.addWidget(self.jumpToLabel)
        Hlayout.addWidget(self.jumpToLabel_2)
        Hlayout.addWidget(self.pageEdit)
        Hlayout.addWidget(self.jumpToLabel_4)
        Hlayout.addWidget(self.pageLabel)
        Hlayout.addWidget(self.jumpToLabel_4)
        Hlayout.addWidget(self.timuBtn)
        self.widget = QWidget()
        self.widget.setLayout(Hlayout)
        self.widget.setFixedWidth(500)

        self.leftPageBtn = QPushButton("上一题")
        self.rightPageBtn = QPushButton("下一题")
        self.leftPageBtn.clicked.connect(self.leftBtnClick)
        self.rightPageBtn.clicked.connect(self.rightBtnClick)
        self.leftPageBtn.setEnabled(False)
        self.rightPageBtn.setEnabled(False)
        # self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionLayout.addLayout(self.optionLayout1)
        self.optionLayout.addLayout(self.optionLayout2)
        # self.optionLayout2.addWidget(self.jumpToLabel_7)
        self.optionPackLayout.addWidget(self.leftPageBtn)
        self.optionPackLayout.addWidget(self.jumpToLabel_7)
        self.optionPackLayout.addLayout(self.optionLayout)
        self.optionPackLayout.addWidget(self.jumpToLabel_7)
        self.optionPackLayout.addWidget(self.rightPageBtn)

        self.optionBtn1.clicked.connect(self.oneBtnClick)
        self.optionBtn2.clicked.connect(self.twoBtnClick)
        self.optionBtn3.clicked.connect(self.threeBtnClick)
        self.optionBtn4.clicked.connect(self.fourBtnClick)
        self.optionBtn5.clicked.connect(self.fiveBtnClick)
        self.optionBtn6.clicked.connect(self.sixBtnClick)
        self.optionBtn7.clicked.connect(self.sevenBtnClick)
        self.optionBtn8.clicked.connect(self.eightBtnClick)
        self.optionBtn9.clicked.connect(self.nineBtnClick)
        self.optionBtn10.clicked.connect(self.tenBtnClick)

        self.exitBtn.setGraphicsEffect(
            QtWidgets.QGraphicsDropShadowEffect(blurRadius=25, xOffset=3, yOffset=3))
        self.sys_manage_btn.setGraphicsEffect(
            QtWidgets.QGraphicsDropShadowEffect(blurRadius=25, xOffset=3, yOffset=3))
        self.studyBtn.setGraphicsEffect(
            QtWidgets.QGraphicsDropShadowEffect(blurRadius=25, xOffset=3, yOffset=3))
        self.user_manage_btn.setGraphicsEffect(
            QtWidgets.QGraphicsDropShadowEffect(blurRadius=25, xOffset=3, yOffset=3))
        self.trainMissionBtn.setGraphicsEffect(
            QtWidgets.QGraphicsDropShadowEffect(blurRadius=25, xOffset=3, yOffset=3))
        self.IndexBtn.setGraphicsEffect(
            QtWidgets.QGraphicsDropShadowEffect(blurRadius=25, xOffset=3, yOffset=3))

        # 视频
        self.videowidget.resize(200, 200)
        self.player.setVideoOutput(self.videowidget)
        self.playLayout = QHBoxLayout()
        self.playpause = QPushButton("播放")
        self.fullplay = QPushButton("全屏播放")
        self.playLayout.addWidget(self.playpause)
        self.playLayout.addWidget(self.fullplay)
        self.video_layout.addWidget(self.videowidget)
        self.video_layout.addLayout(self.playLayout)
        self.video_layout.setStretch(0, 4)
        self.video_layout.setStretch(1, 1)

        # 连词成句
        # 定义1到11的区间(包含1,不包含11),创建w_btn(1-10)赋予w_btn点击事件
        # self.w_btn1= QPushButton("1")
        btn_font=QFont()
        btn_font.setPointSize(15)
        r = range(1, 11)
        for ele in r:
            exec('self.w_btn{}= QPushButton("")'.format(ele))
            exec('self.w_btn{}.clicked.connect(self.wsBtnClick)'.format(ele))
            exec('self.w_btn{}.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)'.format(ele))
            exec('self.w_btn{}.setFont(btn_font)'.format(ele))
            exec('self.s_btn{}.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)'.format(ele))
            exec('self.w_btnArray.append(self.w_btn{})'.format(ele))
            exec('self.s_btnArray.append(self.s_btn{})'.format(ele))
            exec('self.s_btn{}.setFont(btn_font)'.format(ele))
        # 加入排版
        self.ws_layout1.addWidget(self.w_btn1)
        self.ws_layout1.addWidget(self.w_btn2)
        self.ws_layout1.addWidget(self.w_btn3)
        self.ws_layout1.addWidget(self.w_btn4)
        self.ws_layout1.addWidget(self.w_btn5)
        self.ws_layout2.addWidget(self.w_btn6)
        self.ws_layout2.addWidget(self.w_btn7)
        self.ws_layout2.addWidget(self.w_btn8)
        self.ws_layout2.addWidget(self.w_btn9)
        self.ws_layout2.addWidget(self.w_btn10)

        self.ws_layout3.addWidget(self.s_btn1)
        self.ws_layout3.addWidget(self.s_btn2)
        self.ws_layout3.addWidget(self.s_btn3)
        self.ws_layout3.addWidget(self.s_btn4)
        self.ws_layout3.addWidget(self.s_btn5)
        self.ws_layout4.addWidget(self.s_btn6)
        self.ws_layout4.addWidget(self.s_btn7)
        self.ws_layout4.addWidget(self.s_btn8)
        self.ws_layout4.addWidget(self.s_btn9)
        self.ws_layout4.addWidget(self.s_btn10)

        self.checkBtn = QPushButton("确认")
        self.checkBtn.clicked.connect(self.checkClick)
        self.backBtn = QPushButton("回退")
        self.backBtn.clicked.connect(self.minus)

        # 录音相关
        # self.temp_pano = "0"
        self.train_str_no = ""
        self.file_path = ""
        self.trial_no = "0"  # 尚未自动生成
        from myaudiolayout import AudioRecorder
        self.vbox = QHBoxLayout()
        self.audiorecorder = AudioRecorder(self.temp_pano, self.train_str_no, self.trial_no, parent=self)
        self.audiorecorder.pagestate_clicked.connect(self.setBtnsState)  # 改变上下翻页按钮是否激活，防止录音被打断
        self.audiorecorder.hide()  # 去除录音
        self.endRecordBtn = QPushButton("结束录制")  # 手动
        # self.endRecordBtn.setFixedSize(150, 50)
        # self.endRecordBtn.move(250, 100)
        font.setPixelSize(25)
        self.endRecordBtn.setFont(font)
        self.endRecordBtn.clicked.connect(self.endRecordBtnClicked)
        from assess import AudioAssess # version2
        self.audio_assess=AudioAssess()
        self.audio_assess.set_hide()
        self.vbox.addWidget(self.audiorecorder)
        self.vbox.addWidget(self.endRecordBtn)  # 手动
        self.vbox.addWidget(self.audio_assess)

        # 总排版
        self.layout.addLayout(self.indexlayout)
        self.layout.addLayout(self.optionPackLayout)
        self.layout.addLayout(self.contentLayout)
        # 连词成句按钮板块
        label_font=QFont()
        label_font.setPointSize(20)
        self.word_label=QLabel("词语：")
        self.word_label.setFont(label_font)
        self.sentence_label=QLabel("请连词成句：")
        self.sentence_label.setFont(label_font)

        self.ws_layout.addWidget(self.word_label)
        self.ws_layout.addLayout(self.ws_layout1)
        self.ws_layout.addLayout(self.ws_layout2)
        self.ws_layout.addWidget(self.sentence_label)
        self.ws_layout.addLayout(self.ws_layout3)
        self.ws_layout.addLayout(self.ws_layout4)
        # 确认和回退按钮
        self.checkpack_layout = QHBoxLayout()
        self.checkpack_layout.addWidget(self.checkBtn)
        self.checkpack_layout.addWidget(self.backBtn)

        # 连词成句按钮板块排版
        self.right_layout=QVBoxLayout()
        self.right_layout.addLayout(self.ws_layout,8)
        self.right_layout.addLayout(self.checkpack_layout,1)
        self.right_layout.addLayout(self.vbox,1)

        self.contentLayout.addLayout(self.video_layout, 2)
        self.contentLayout.addLayout(self.right_layout, 3)
        self.playpause.clicked.connect(self.openVideoFile)
        self.fullplay.clicked.connect(self.fullopenVedio)
        self.layout.addWidget(self.widget, 1, Qt.AlignCenter)

        self.setLayout(self.layout)

        lab = [self.endRecordBtn]
        tb = [self.exitBtn, self.sys_manage_btn, self.studyBtn, self.user_manage_btn, self.IndexBtn
              ]  # ,
        yb = [self.optionBtn1, self.optionBtn2, self.optionBtn3, self.optionBtn4, self.optionBtn5, self.optionBtn6,
              self.optionBtn7, self.optionBtn8, self.optionBtn9, self.optionBtn10,
              ]
        lb = [self.leftPageBtn, self.rightPageBtn]
        hb = [self.playpause, self.fullplay]
        # qb = [self.vowelBtn]
        font = QtGui.QFont()
        font.setPointSize(15)  # 括号里的数字可以设置成自己想要的字体大小
        # font.setFamily("SimHei")  # 黑体
        font.setFamily("SimSun")  # 宋体
        for i in lab:
            i.setFixedSize(150, 35)
            i.setStyleSheet("QPushButton{\n"
                            "    background-color: rgb(193, 193, 193);\n"
                            "    color: rgb(81, 71, 81);\n"
                            "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 5px;font-family: 微软雅黑;\n"
                            "}\n"
                            "QPushButton:pressed{\n"
                            "    background:black;\n"
                            "}")
            i.setFont(font)
        for i in tb:
            i.setStyleSheet("QPushButton{\n"
                            "    background:rgb(244, 183, 0);\n"
                            "    color: rgb(81, 71, 81);\n"
                            "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 24px;font-family: 微软雅黑;\n"
                            "}\n"
                            "QPushButton:pressed{\n"
                            "    background:black;\n"
                            "}")
            i.setFont(font)
        for i in yb:
            i.setFixedSize(150, 30)
            i.setStyleSheet("QPushButton{\n"
                            "    background-color: rgb(170, 255, 127);\n"
                            "    color: rgb(81, 71, 81);\n"
                            "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 5px;font-family: 微软雅黑;\n"
                            "}\n"
                            "QPushButton:pressed{\n"
                            "    background:yellow;\n"
                            "}")
            i.setFont(font)
            font = QtGui.QFont()
            font.setPointSize(10)  # 括号里的数字可以设置成自己想要的字体大小
            i.setFont(font)
        for i in lb:
            i.setFixedSize(450, 35)
            i.setStyleSheet("QPushButton{\n"
                            "    background-color: rgb(170, 255, 127);\n"
                            "    color: rgb(81, 71, 81);\n"
                            "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 5px;font-family: 微软雅黑;\n"
                            "}\n"
                            "QPushButton:pressed{\n"
                            "    background:yellow;\n"
                            "}")
            i.setFont(font)
            font = QtGui.QFont()
            font.setPointSize(10)  # 括号里的数字可以设置成自己想要的字体大小
            i.setFont(font)
        for i in hb:
            i.setFixedSize(400, 35)
            i.setStyleSheet("QPushButton{\n"
                            "    background-color: rgb(170, 255, 127);\n"
                            "    color: rgb(81, 71, 81);\n"
                            "    box-shadow: 1px 1px 3px;font-size:18px;border-radius: 5px;font-family: 微软雅黑;\n"
                            "}\n"
                            "QPushButton:pressed{\n"
                            "    background:yellow;\n"
                            "}")
            i.setFont(font)
            font = QtGui.QFont()
            font.setPointSize(10)  # 括号里的数字可以设置成自己想要的字体大小
            i.setFont(font)

        # 录音相关，在录音时控制翻页是否可用

    def userManageBtnClicked(self):
        self.userManageWindow.showFullScreen()
        self.userManageWindow.searchButtonClicked()
        self.close()

    def studyBtnClicked(self):
        QtGui.QDesktopServices.openUrl(QtCore.QUrl('https://book.yunzhan365.com/vdogo/wpil/mobile/index.html'))

    # 录音相关，在录音时控制翻页是否可用
    def setBtnsState(self, b_value):
        self.leftPageBtn.setEnabled(b_value)
        self.rightPageBtn.setEnabled(b_value)
        self.endRecordBtn.setEnabled(not b_value)

    def IndexBtnClick(self):
        from t_main import IndexWindow
        self.mainwindow = IndexWindow()
        self.mainwindow.showFullScreen()
        self.mainwindow.searchButtonClicked()
        self.close()

    def trainMissionBtnClick(self):
        if not (self.temp_pano == ""):
            from train import TrainWindow
            self.trainWindow = TrainWindow()
            self.trainWindow.showFullScreen()
            self.trainWindow.setNo(self.temp_pano)
            self.close()
        else:
            print(QMessageBox.warning(self, "警告", "请选择一名患者", QMessageBox.Yes, QMessageBox.Yes))

    def AssignSentence(self):
        db = QSqlDatabase.addDatabase("QSQLITE")
        db.setDatabaseName('database.db')
        db.open()
        query = QSqlQuery()
        sql = "SELECT COUNT(DISTINCT sentence_no) as num_common_sno FROM sentence WHERE sentence_no "+\
              "IN (SELECT sentence_no FROM sentence_word)"
        query.exec_(sql)
        query.next()
        sentence_in_sql = int(query.value(0))
        if (sentence_in_sql < self.sentence_total):
            temp_num = sentence_in_sql
            self.sentence_total = temp_num
        else:
            temp_num = self.sentence_total
        if self.is_random==0:
            sql = "SELECT DISTINCT sentence.* FROM sentence INNER JOIN sentence_word"\
                  +" ON sentence.sentence_no = sentence_word.sentence_no ORDER BY sentence_no LIMIT '%s'" % (temp_num)
        else:
            sql = "SELECT DISTINCT sentence.* FROM sentence INNER JOIN sentence_word" \
                  + " ON sentence.sentence_no = sentence_word.sentence_no ORDER BY RANDOM() LIMIT '%s'" % (temp_num)
        query.exec_(sql)
        r = range(0, temp_num)
        for ele in r:
            query.next()
            self.sArray.append([])
            print(query.value(0))
            self.sArray[ele].append(query.value(0))  # sentence_no
            self.sArray[ele].append(query.value(1))  # sentence_name
            self.sArray[ele].append(query.value(2))  # sentence_audio
        # 训练录制相关
        # 定义查询语句和插入语句
        select_sql = "SELECT trial_no FROM trial WHERE pa_no = ?"
        insert_sql = "INSERT INTO trial (pa_no, trial_no) VALUES (?, 0)"
        update_sql = "UPDATE trial SET trial_no = ? WHERE pa_no = ?"
        # 定义要查询的pa_no值
        pa_no = self.temp_pano
        # 执行查询操作
        query = QSqlQuery()
        query.prepare(select_sql)
        query.addBindValue(pa_no)
        query.exec_()
        if not query.next():
            # 如果查询结果为空，则插入新记录
            query.prepare(insert_sql)
            query.addBindValue(pa_no)
            query.exec_()
            db.commit()
            self.trial_no = 0
        else:
            # 如果查询结果不为空，则更新记录
            trial_no = query.value(0) + 1
            query.prepare(update_sql)
            query.addBindValue(trial_no)
            query.addBindValue(pa_no)
            query.exec_()
            db.commit()
            self.trial_no = trial_no

    def AssignBtns(self):
        # 录音相关
        self.sentence_no = self.sArray[self.current_s][0]
        self.train_str_no = "s" + str(self.sentence_no)
        self.audiorecorder.setTrainStr(self.train_str_no)
        self.audiorecorder.setTrialNo(self.trial_no)
        self.audiorecorder.setFilePath(self.trial_no, self.train_str_no)
        self.audiorecorder.setAuBtnsState(False)
        self.rightPageBtn.setEnabled(False)

        self.i = 1
        temp_s_no = self.sArray[self.current_s][0]
        db = QSqlDatabase.addDatabase("QSQLITE")
        db.setDatabaseName('database.db')
        db.open()
        query = QSqlQuery()
        sql = "SELECT word_name FROM sentence_word WHERE sentence_no= '%s'" % (temp_s_no)
        query.exec_(sql)
        # 添加词语到列表
        while (query.next()):
            self.wArray.append(query.value(0))
        random.shuffle(self.wArray)
        # 设置词语按钮
        temp_w_count = len(self.wArray)
        r = range(0, temp_w_count)
        for ele in r:
            temp_str = self.wArray[ele]
            self.w_btnArray[ele].setText(temp_str)
            self.w_btnArray[ele].setEnabled(True)
            self.s_btnArray[ele].setEnabled(True)
        r = range(temp_w_count, 10)
        for ele in r:
            self.w_btnArray[ele].setText("")
            self.w_btnArray[ele].setEnabled(False)
            self.s_btnArray[ele].setEnabled(False)
        r = range(0, 10)
        for ele in r:
            self.s_btnArray[ele].setText("")
        # 设置播放
        self.media_path = "pa_video/pa" + str(self.temp_pano)+ "/pa" + str(self.temp_pano) + "_s" + str(temp_s_no) + ".mp4"
        self.player.setMedia(QMediaContent(QUrl.fromLocalFile(self.media_path)))
        self.openVideoFile()

    def openVideoFile(self):
        self.player.setVolume(80)
        self.player.play()

    def fullopenVedio(self):
        self.fsDialog.setMedia(self.media_path)
        self.fsDialog.showFullScreen()
        self.fsDialog.play()
        return

    def timuBtnClick(self):
        if (self.pageEdit.text().isdigit()):
            print(self.sentence_total)
            s = "/" + str(int(self.sentence_total)) + "题"
            self.pageLabel.setText(s)
            self.currentPage = int(self.pageEdit.text())
            if (self.currentPage > self.sentence_total):
                # self.currentPage = self.vowel_total
                es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
                if (es == QMessageBox.Yes):
                    print("ok")
                else:
                    return
            if (self.currentPage < 1):
                # self.currentPage = 1
                es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
                if (es == QMessageBox.Yes):
                    print("ok")
                else:
                    return
            else:
                self.current_s = self.currentPage - 1
                self.AssignBtns()
        else:
            es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
            if (es == QMessageBox.Yes):
                print("ok")
            else:
                return
    def oneBtnClick(self):
        print("总数total:")
        print(self.vowel_total)
        print(self.current_v)
        if (self.current_s >= 0 and self.sentence_total >= 1):#and self.word_total>=1
            self.current_s = 0
            self.wArray.clear()
            self.AssignBtns()
        else:
          es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
          if (es == QMessageBox.Yes):
            print("ok")
          else:
            return

    def twoBtnClick(self):
        if (self.current_s >= 0 and self.sentence_total>=2):#and self.word_total>=1
            self.current_s = 1
            self.wArray.clear()
            self.AssignBtns()
        else:
           es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
           if (es == QMessageBox.Yes):
             print("ok")
           else:
             return
    def threeBtnClick(self):
        if (self.current_s >= 0 and self.sentence_total>=3):#and self.word_total>=1
            self.current_s = 2
            self.wArray.clear()
            self.AssignBtns()
        else:
           es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
           if (es == QMessageBox.Yes):
             print("ok")
           else:
             return
    def fourBtnClick(self):
        if (self.current_s >= 0 and self.sentence_total>=4):#and self.word_total>=1
            self.current_s = 3
            self.wArray.clear()
            self.AssignBtns()
        else:
           es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
           if (es == QMessageBox.Yes):
             print("ok")
           else:
             return

    def fiveBtnClick(self):
        if (self.current_s >= 0 and self.sentence_total >= 5):  # and self.word_total>=1
            self.current_s = 4
            self.wArray.clear()
            self.AssignBtns()
        else:
            es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
            if (es == QMessageBox.Yes):
                print("ok")
            else:
                return
    def sixBtnClick(self):
        if (self.current_s >= 0 and self.sentence_total >= 6):  # and self.word_total>=1
            self.current_s = 5
            self.wArray.clear()
            self.AssignBtns()
        else:
            es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
            if (es == QMessageBox.Yes):
                print("ok")
            else:
                return
    def sevenBtnClick(self):
        if (self.current_s >= 0 and self.sentence_total >= 7):  # and self.word_total>=1
            self.current_s = 6
            self.wArray.clear()
            self.AssignBtns()
        else:
            es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
            if (es == QMessageBox.Yes):
                print("ok")
            else:
                return
    def eightBtnClick(self):
        if (self.current_s >= 0 and self.sentence_total >= 8):  # and self.word_total>=1
            self.current_s = 7
            self.wArray.clear()
            self.AssignBtns()
        else:
            es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
            if (es == QMessageBox.Yes):
                print("ok")
            else:
                return
    def nineBtnClick(self):
        if (self.current_s >= 0 and self.sentence_total >= 9):  # and self.word_total>=1
            self.current_s = 8
            self.wArray.clear()
            self.AssignBtns()
        else:
            es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
            if (es == QMessageBox.Yes):
                print("ok")
            else:
                return

    def tenBtnClick(self):
        if (self.current_s >= 0 and self.sentence_total >= 10):  # and self.word_total>=1
            self.current_s = 9
            self.wArray.clear()
            self.AssignBtns()
        else:
            es = QMessageBox.information(self, "提示", "没有此题，请返回", QMessageBox.Yes)
            if (es == QMessageBox.Yes):
                print("ok")
            else:
                return

    def leftBtnClick(self):
        if (self.current_s > 0):
            self.current_s -= 1
            self.wArray.clear()
            self.AssignBtns()

    def rightBtnClick(self):
        self.playflag = 0  # 训练录制相关
        self.read_no=0 # version2
        self.stop_timer()
        if (self.current_s < self.sentence_total - 1):
            self.current_s += 1
            self.wArray.clear()
            self.AssignBtns()
            self.merge_html_files()
            print("go on")
        else:
            self.recording_thread.set_close()  # 相机逻辑
            self.writelogs()
            ress = QMessageBox.information(self, "提示", "是否查看本次训练结果?", QMessageBox.Yes, QMessageBox.No)
            if (ress == QMessageBox.Yes):
                self.writetubiao()
            else:
                return
            #self.wt, self.rt, self.nt,self.wrongword,self.rightword
            # 总结报告
            from train_log import TrainDialogue
            traindialogue=TrainDialogue(self)
            traindialogue.exec_()
            self.trainMissionBtnClick()
    def merge_html_files(self):
        target_file_path = self.temp_pano + 'target_merged.html'  # 假设target_merged.html应该在工作目录中

        # 检查target_merged.html文件是否存在，如果不存在则创建
        if not os.path.exists(target_file_path):
            with open(target_file_path, 'w', encoding='utf-8') as file:
                file.write('<html><body></body></html>')  # 创建一个简单的HTML结构

        # 读取target_merged.html文件的内容
        with open(target_file_path, 'r', encoding='utf-8') as file:
            target_content = file.read()

        # 在target_merged.html文件中找到合适的位置插入源HTML文件的内容
        insert_position = target_content.find('</body>')
        if insert_position != -1:
            # 插入所有源HTML文件的内容
            for source_file in self.source_files:
                if os.path.exists(source_file):
                    with open(source_file, 'r', encoding='utf-8') as file:
                        source_content = file.read()
                        target_content = target_content[:insert_position] + source_content + target_content[
                                                                                             insert_position:]
                else:
                    print(f"File {source_file} does not exist.")
        else:
            # 如果没有找到</body>标签，可以选择在末尾添加或者进行其他操作
            for source_file in self.source_files:
                if os.path.exists(source_file):
                    with open(source_file, 'r', encoding='utf-8') as file:
                        source_content = file.read()
                        target_content += source_content
                else:
                    print(f"File {source_file} does not exist.")

        # 将新的内容写回target_merged.html文件
        with open(target_file_path, 'w', encoding='utf-8') as file:
            file.write(target_content)
    def writetubiao(self):
        from pyecharts import options as opts
        from pyecharts.charts import Bar
        now = datetime.datetime.now()
        # 将日期格式化为SQLite可接受的格式
        log_date = now.strftime('%Y-%m-%d')
        c = (
            Bar()
            .add_xaxis(
                [
                    log_date
                ]
            )
            .add_yaxis("错题数", [self.wt])
            # .add_yaxis("普通读音数", [self.nt])
            .add_yaxis("正确数", [self.rt])
            .set_global_opts(
                # xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=-15)),  # 此行代码旋转标签的角度
                title_opts=opts.TitleOpts(title="患者" + self.temp_pano + "连词成句训练"),
            )
            .render(self.temp_pano+"lc.html")
        )
        import webbrowser
        webbrowser.open(self.temp_pano+'lc.html')

    def writealltubiao(self):
        from pyecharts import options as opts
        from pyecharts.charts import Bar

        now = datetime.datetime.now()
        # 将日期格式化为SQLite可接受的格式
        log_date = now.strftime('%Y-%m-%d')
        c = (
            Bar()
            .add_xaxis(
                [
                    "第" + self.trial_no + "次训练"
                ]
            )
            .add_yaxis("连词成句错题数", [self.wrongword])
            .add_yaxis("连词成句正确数", [self.rightword])
            .set_global_opts(
                # xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=-15)),  # 此行代码旋转标签的角度
                title_opts=opts.TitleOpts(title="患者" + self.temp_pano,
                                          subtitle=log_date),
            )
            .render("label.html")
        )
        import webbrowser
        webbrowser.open('label.html')

    def writelogs(self):
        # 连接到数据库
        conn = sqlite3.connect('database.db')
        # 创建游标
        cursor = conn.cursor()
        # 查询是否存在匹配的记录
        query = '''
               SELECT * FROM logs WHERE pa_no = ? AND log_date = ?
           '''
        # 获取当前日期
        now = datetime.datetime.now()
        # 将日期格式化为SQLite可接受的格式
        log_date = now.strftime('%Y-%m-%d')
        cursor.execute(query, (self.temp_pano, log_date))
        record = cursor.fetchone()
        # 如果不存在记录，则插入新记录
        if not record:
            # 计算acc的值
            acc = 100 - int(self.wrong_num / self.sum_num * 100)
            # 插入新记录
            query = '''
                   INSERT INTO logs (pa_no, log_date, wrong_num, sum_num, acc)
                   VALUES (?, ?, ?, ?, ?)
               '''
            values = (self.temp_pano, log_date, self.wrong_num, self.sum_num, acc)
            cursor.execute(query, values)
        else:
            # 计算更新后的wrong_num、sum_num和acc的值
            wrong_num = record[3] + self.wrong_num
            sum_num = record[4] + self.sum_num
            acc = 100 - int(wrong_num / sum_num * 100)
            # 更新记录
            query = '''
                   UPDATE logs SET wrong_num = ?, sum_num = ?, acc = ?
                   WHERE pa_no = ? AND log_date = ?
               '''
            values = (wrong_num, sum_num, acc, self.temp_pano, log_date)
            cursor.execute(query, values)
        # 提交更改并关闭游标和连接
        conn.commit()
        cursor.close()
        conn.close()
        # 创建docx文档
        doc = docx.Document()
        p = doc.add_paragraph()
        p.add_run(log_date + '\n')
        # 添加窗口文本到文档中
        for i, (ques, score, text) in enumerate(zip(self.ques_list, self.score_list, self.text_list)):
            # 将窗口文本添加到docx文档中
            p.add_run(ques + '\n').bold = True  # 加粗显示问题
            p.add_run(score + '\n')  # 显示分数
            # 使用 MyHTMLParser 解析 HTML 文本
            parser = MyHTMLParser()
            parser.feed(text)
            # 处理文本和属性
            for parsed_text, bold, color in parser.result:
                run = p.add_run(parsed_text)
                # 处理加粗
                if bold:
                    run.bold = True
                # 处理颜色
                if color:
                    run.font.color.rgb = docx.shared.RGBColor(*color)
        filename = 'pa_train/pa_' + str(self.temp_pano) + '/trial' + str(self.trial_no) + '/output.docx'
        # 保存为docx文件
        doc.save(filename)

        filename = 'pa_train/pa_' + str(self.temp_pano) + '/output.docx'
        if os.path.exists(filename):
            # 如果文件已经存在，则打开现有文件以追加模式
            doc = docx.Document(filename)
        else:
            # 如果文件不存在，则创建一个新文件
            doc = docx.Document()
        p = doc.add_paragraph()
        p.add_run(log_date + '\n')
        p.add_run('trial' + str(self.trial_no) + '\n')
        # 在现有文档或新文档中添加窗口文本
        for i, (ques, score, text) in enumerate(zip(self.ques_list, self.score_list, self.text_list)):
            # 将窗口文本添加到docx文档中
            p = doc.add_paragraph()
            p.add_run(ques + '\n').bold = True  # 加粗显示问题
            p.add_run(score + '\n')  # 显示分数
            # 使用 MyHTMLParser 解析 HTML 文本
            parser = MyHTMLParser()
            parser.feed(text)
            # 处理文本和属性
            for parsed_text, bold, color in parser.result:
                run = p.add_run(parsed_text)
                # 处理加粗
                if bold:
                    run.bold = True
                # 处理颜色
                if color:
                    run.font.color.rgb = docx.shared.RGBColor(*color)
        # 保存为docx文件
        doc.save(filename)

    def wsBtnClick(self):
        sender = self.sender()
        senderText = str(sender.text())
        if not senderText == "" and self.i <= len(self.wArray):
            if self.countflag == 2:
                self.i += 1
            exec('self.s_btn{}.setText(senderText)'.format(self.i))
            audio_path = "audio/word/" + senderText + ".mp3"
            self.audio_player.setMedia(QMediaContent(QUrl.fromLocalFile(audio_path)))
            self.audio_player.setVolume(80)
            self.audio_player.play()
            if self.i >= 0:
                self.i += 1  # 变量自增
                self.countflag = 1
        else:
            return

    def minus(self):
        print("i的值是" + str(self.i))
        print("array的长度为" + str(len(self.wArray)))
        exec('self.s_btn{}.setText("")'.format(self.i))
        if (self.countflag == 1):
            exec('self.s_btn{}.setText("")'.format(self.i))
            self.i -= 1
            self.countflag = 0
        exec('self.s_btn{}.setText("")'.format(self.i))
        if self.i > 1:
            self.i -= 1
        if not self.s_btn1.text() == "":  # 减了，但是没减到底
            self.countflag = 2
        else:
            self.countflag = 0
        print("coutflag=" + str(self.countflag))

    def checkClick(self):
        temp_sentence = ""
        r = range(0, 10)
        for ele in r:
            temp_sentence = temp_sentence + self.s_btnArray[ele].text()
        if temp_sentence == self.sArray[self.current_s][1]:
            print({QMessageBox.information(self, "提示", "回答正确！", QMessageBox.Yes, QMessageBox.Yes)})
            self.rt += 1
            print("self.rt")
            print(self.rt)
            self.rightBtnClick()
        else:
            print({QMessageBox.information(self, "提示", "回答错误！再试一遍吧！", QMessageBox.Yes, QMessageBox.Yes)})
            self.wt += 1
            print("self.wt")
            print(self.wt)

    def exitBtnClick(self):
        ret = QMessageBox.information(self, "提示", "是否退出系统?", QMessageBox.Yes, QMessageBox.No)
        if (ret == QMessageBox.Yes):
            sys.exit(app.exec_())
        else:
            return

    def sysManageBtnClicked(self):
        from sys_management import SysManageWindow
        self.sysManageWindow=SysManageWindow()
        self.sysManageWindow.searchButtonClicked()
        self.sysManageWindow.showFullScreen()
        self.close()

    def setTaskNum(self, pano_value, num_value,is_random,is_auto,repet_no):
        self.temp_pano = pano_value
        self.sentence_total = num_value
        self.is_random = is_random
        self.is_auto = is_auto
        if repet_no>=1:
            self.repet_no=repet_no

    def setTrainFile(self):#训练录制相关
        self.trainfolder="pa_train/pa_"+str(self.temp_pano)+"/trial"+str(self.trial_no)+"/"
        if not os.path.exists(self.trainfolder):
            os.makedirs(self.trainfolder)
        self.trainfile=self.trainfolder+self.train_str_no+"_"+str(self.read_no)+".mp4" #version2

    def on_state_changed(self, state):
        try:
            if self.is_auto == 0:
                if state == QMediaPlayer.StoppedState and self.playflag<=self.repet_no:
                    self.read_no=self.read_no+1 # version2
                    self.duration = (self.player.duration() / 1000 * 2 + 6)
                    # 训练录制相关
                    self.setTrainFile()
                    print("录制中"+str(self.current_s))
                    self.recording_thread.set_parameter(self.duration, self.trainfile)  # 相机逻辑
                    self.recording_thread.start()
                    self.setBtnsState(False)
                    self.msgBox.setText("请跟读!!")
                    self.msgtimer.start(700)  # 提示
                    self.msgBox.exec_()
                    self.playflag = self.playflag+1
                    time.sleep(self.duration)
                    temp_standard_text=self.sArray[self.current_s][1]
                    temp_standard_audio="audio/sentence/"+str(self.sentence_no)+".mp3"
                    temp_user_audio="user_audio.mp3"
                    # 总结报告
                    score,text, wrong_num, sum_num=self.audio_assess.set_file(temp_standard_text,temp_standard_audio,temp_user_audio)

                    os.remove("user_audio.mp3")
                    ques="第"+str(self.read_no)+"遍："+temp_standard_text
                    self.record_score_and_text(ques,score,text)
                    self.wrong_num += wrong_num
                    self.sum_num += sum_num
                    # self.sc = self.audio_assess.sc
                    # print("这是词语界面的" + self.sc)
                    # # from assess import AudioAssess
                    # if int(self.sc) > 85:
                    #     self.rt += 1
                    # elif int(self.sc) <= 85 and int(self.sc) >= 70:
                    #     self.nt += 1
                    # else:
                    #     self.wt += 1
                    time.sleep(1)
                    if self.playflag==self.repet_no:
                        self.msgBox.setText("跟读结束，请根据刚才的句子进行连词成句")
                        self.msgtimer.start(1000)  # 提示
                        self.msgBox.exec_()
                        self.timer_thread.start()
                        self.timer_thread.change_interval(int(self.duration))  # 计时器，计时结束则自动翻页
                    else:
                        self.openVideoFile()
            # 手动
            elif self.is_auto == 1:
                if state == QMediaPlayer.StoppedState:
                    self.read_no = self.read_no + 1  # version2
                    # 训练录制相关
                    self.setTrainFile()
                    print("录制中" + str(self.current_s))
                    self.recording_thread.set_parameter(0, self.trainfile)  # 相机逻辑
                    self.recording_thread.start()
                    self.setBtnsState(False)
                    self.msgBox.setText("按 结束录制 结束跟读")
                    self.msgtimer.start(700)  # 提示
                    # self.sc = self.audio_assess.sc
                    # print("这是词语界面的" + self.sc)
                    # # from assess import AudioAssess
                    # if int(self.sc) > 85:
                    #     self.rt += 1
                    # elif int(self.sc) <= 85 and int(self.sc) >= 70:
                    #     self.nt += 1
                    # else:
                    #     self.wt += 1
                    self.msgBox.exec_()
                    self.playflag = self.playflag + 1
        except Exception as e:
            print(f'Error: {e}')

    # 手动
    def endRecordBtnClicked(self):
        try:
            self.recording_thread.set_vr_flag(True)
            time.sleep(1)
            temp_standard_audio = "audio/sentence/" + str(self.sentence_no) + ".mp3"
            temp_user_audio = "user_audio.mp3"
            temp_standard_text = self.sArray[self.current_s][1]
            # 总结报告
            score, text, wrong_num, sum_num = self.audio_assess.set_file(temp_standard_text, temp_standard_audio, temp_user_audio)
            self.sum_num += sum_num
            self.wrong_num += wrong_num
            os.remove("user_audio.mp3")
            ques = "第" + str(self.read_no) + "遍：" + temp_standard_text
            self.record_score_and_text(ques, score, text)
            time.sleep(1)
            self.setBtnsState(True)
            self.recording_thread.set_vr_flag(False)
        except Exception as e:
            print(f'Error: {e}')

    # 总结报告
    def record_score_and_text(self,ques, score, text):
        self.ques_list.append(ques)
        self.score_list.append(score)
        self.text_list.append(text)

    #计时器相关
    def timer_callback(self):
        self.rightBtnClick()

    @pyqtSlot()
    def start_timer(self):
        self.timer_thread.start()

    @pyqtSlot()
    def stop_timer(self):
        self.timer_thread.stop()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon("./images/MainWindow_1.png"))
    mainWindow = WSWindow()
    temp_pano = 0
    mainWindow.setTaskNum(temp_pano, int(5))
    #is_random=1,is_auto=1,repet_no=1
    mainWindow.AssignSentence()
    mainWindow.AssignBtns()
    mainWindow.showFullScreen()
    sys.exit(app.exec_())